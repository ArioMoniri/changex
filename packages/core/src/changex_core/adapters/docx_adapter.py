"""The docx adapter: load -> model -> apply v0.1 ops -> render native revisions.

Identity strategy
-----------------
* Each docx paragraph reuses Word's native ``w14:paraId``. Files that lack it
  (some generators omit paraIds) get one **minted and injected** on load, so the
  id is stable across a save/reload round-trip.
* Paragraph node_ids therefore survive text edits and sibling insertions — the
  whole point of decoupling identity from content.

Revision model
--------------
The adapter tracks, per paragraph, an ordered list of text **segments** tagged
``keep`` / ``ins`` / ``del``. Applying ops splits/retags segments; rendering
turns segments into ``w:ins``/``w:del``(delText)/plain runs. Consequently:

* **accept-all** drops every ``del`` segment and unwraps every ``ins`` → the
  active-capture target text.
* **reject-all** drops every ``ins`` segment and unwraps every ``del`` → the
  baseline text.

These two equalities are the gating M0 round-trip criterion.
"""

from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn

from changex_core.adapters.base import (
    BeforeMismatchError,
    DocumentAdapter,
    NodeNotFoundError,
    OversizedOpError,
)
from changex_core.adapters.docx_revisions import (
    WidAllocator,
    build_run_props,
    make_deleted_pilcrow,
    make_deleted_run,
    make_inserted_pilcrow,
    make_inserted_run,
    make_ppr_change,
    make_rpr_change,
)
from changex_core.journal.canonical import sha256_hex
from changex_core.model.addressing import NodeIdAllocator
from changex_core.model.nodes import Node, NodeKind
from changex_core.ops.vocabulary import (
    FormatRun,
    NodeDelete,
    NodeInsert,
    NodeMove,
    Op,
    StyleChange,
    TextDelete,
    TextInsert,
    TextReplace,
)
from changex_core.paths import safe_path

# An op may not rewrite more than this fraction of a node's text in one go.
MAX_OP_FRACTION = 0.5
DEFAULT_AUTHOR = "ChangeX agent"
DEFAULT_DATE = "2026-06-02T00:00:00Z"


@dataclass
class _Segment:
    """One contiguous text segment within a paragraph and its revision tag."""

    text: str
    tag: str  # "keep" | "ins" | "del"


@dataclass
class _Para:
    """Adapter-side state for one paragraph (model + revision segments)."""

    node_id: str
    para_id: str
    style: str
    base_style: str
    style_changed: bool = False
    segments: list[_Segment] = field(default_factory=list)
    inserted: bool = False  # whole paragraph is a node.insert
    deleted: bool = False  # whole paragraph is a node.delete
    # format.run state: the live run props + the prior props (for w:rPrChange).
    format_changed: bool = False
    format_props: dict = field(default_factory=dict)
    format_before: dict = field(default_factory=dict)

    def current_text(self) -> str:
        """Text as it stands now (keep + ins segments; del removed)."""
        return "".join(s.text for s in self.segments if s.tag != "del")

    def baseline_text(self) -> str:
        """Text in the baseline (keep + del segments; ins removed)."""
        return "".join(s.text for s in self.segments if s.tag != "ins")


class DocxAdapter(DocumentAdapter):
    """Loads a .docx, applies v0.1 ops, renders native Word revisions."""

    def __init__(
        self,
        document: "Document",
        raw_bytes: bytes,
        *,
        author: str = DEFAULT_AUTHOR,
        date: str = DEFAULT_DATE,
    ) -> None:
        self._doc = document
        self._raw = raw_bytes
        self._author = author
        self._date = date
        self._baseline_sha = sha256_hex(raw_bytes)
        self._allocator = NodeIdAllocator()
        self._paras: list[_Para] = []
        self._build_model()

    # -- construction ---------------------------------------------------------

    @classmethod
    def load(cls, path: str, *, author: str = DEFAULT_AUTHOR, date: str = DEFAULT_DATE) -> "DocxAdapter":
        """Load a .docx from a sanitized path."""
        resolved = safe_path(path, must_exist=True, allow_suffixes=(".docx",))
        raw = resolved.read_bytes()
        document = Document(str(resolved))
        return cls(document, raw, author=author, date=date)

    def _build_model(self) -> None:
        self._paras = []
        for p in self._doc.paragraphs:
            para_id = self._ensure_para_id(p)
            node_id = self._allocator.for_para_id(para_id)
            style = p.style.name if p.style is not None else "Normal"
            text = p.text
            segments = [_Segment(text=text, tag="keep")] if text else []
            self._paras.append(
                _Para(
                    node_id=node_id,
                    para_id=para_id,
                    style=style,
                    base_style=style,
                    segments=segments,
                )
            )

    def _ensure_para_id(self, paragraph) -> str:  # type: ignore[no-untyped-def]
        """Return the paragraph's w14:paraId, minting+injecting one if absent."""
        existing = paragraph._p.get(qn("w14:paraId"))
        if existing:
            return str(existing).upper()
        minted = f"{(len(self._paras) + 1):08X}"
        paragraph._p.set(qn("w14:paraId"), minted)
        return minted

    # -- DocumentAdapter contract --------------------------------------------

    def baseline_sha256(self) -> str:
        return self._baseline_sha

    def to_model(self) -> Node:
        """Return the current model tree (paragraphs with a single text run)."""
        root = Node(node_id="root", node_kind=NodeKind.DOCUMENT, path="/body")
        for idx, para in enumerate(self._paras):
            if para.deleted:
                continue
            pnode = Node(
                node_id=para.node_id,
                node_kind=NodeKind.PARAGRAPH,
                path=f"/body/p[{idx + 1}]",
                value=para.current_text(),
                attrs={"style": para.style, "para_id": para.para_id},
            )
            root.children.append(pnode)
        return root

    def set_model(self, root: Node) -> None:
        """Reset adapter state to ``root`` (used by replay onto a baseline)."""
        self._paras = []
        self._allocator = NodeIdAllocator()
        for idx, pnode in enumerate(root.child_paragraphs()):
            para_id = str(pnode.attrs.get("para_id", f"{idx + 1:08X}"))
            self._allocator.for_para_id(para_id)
            text = str(pnode.value or "")
            style = str(pnode.attrs.get("style", "Normal"))
            self._paras.append(
                _Para(
                    node_id=pnode.node_id,
                    para_id=para_id,
                    style=style,
                    base_style=style,
                    segments=[_Segment(text=text, tag="keep")] if text else [],
                )
            )

    def resolve(self, node_id: str) -> Node | None:
        return self.to_model().find(node_id)

    def _para(self, node_id: str) -> _Para:
        for para in self._paras:
            if para.node_id == node_id and not para.deleted:
                return para
        raise NodeNotFoundError(f"no paragraph with node_id {node_id!r}")

    def _para_index(self, node_id: str) -> int:
        for i, para in enumerate(self._paras):
            if para.node_id == node_id:
                return i
        raise NodeNotFoundError(f"no paragraph with node_id {node_id!r}")

    # -- apply ----------------------------------------------------------------

    def apply(self, op: Op) -> None:
        """Apply one v0.1 op (validating before-substring and op size)."""
        if isinstance(op, TextReplace):
            self._apply_replace(op)
        elif isinstance(op, TextInsert):
            self._apply_insert(op)
        elif isinstance(op, TextDelete):
            self._apply_delete(op)
        elif isinstance(op, StyleChange):
            self._apply_style(op)
        elif isinstance(op, NodeInsert):
            self._apply_node_insert(op)
        elif isinstance(op, NodeDelete):
            self._apply_node_delete(op)
        elif isinstance(op, FormatRun):
            self._apply_format_run(op)
        elif isinstance(op, NodeMove):
            self._apply_node_move(op)
        else:  # pragma: no cover - exhaustive over the frozen set
            raise TypeError(f"unsupported op type {type(op).__name__}")

    def _check_size(self, para: _Para, changed_len: int) -> None:
        current = len(para.current_text())
        if current and changed_len > current * MAX_OP_FRACTION:
            raise OversizedOpError(
                "split_required: this op rewrites more than "
                f"{int(MAX_OP_FRACTION * 100)}% of the node; split it into "
                "smaller edits and re-issue them."
            )

    def _locate(self, para: _Para, needle: str) -> tuple[int, int]:
        """Return (segment_index, offset) where ``needle`` starts in a 'keep' run.

        Searches the current (keep+ins) text but resolves to a single keepable
        segment so we can split it. Raises on mismatch (the before-validation).
        """
        if needle == "":
            raise BeforeMismatchError("before/anchor must be non-empty")
        cursor = 0
        for i, seg in enumerate(para.segments):
            if seg.tag == "del":
                continue
            start = seg.text.find(needle)
            if start != -1:
                return i, start
            cursor += len(seg.text)
        raise BeforeMismatchError(
            f"before text {needle!r} not found in node {para.node_id!r}"
        )

    def _split_segment(self, para: _Para, seg_idx: int, start: int, length: int) -> int:
        """Split the segment at ``seg_idx`` so [start, start+length) is isolated.

        Returns the index of the isolated middle segment.
        """
        seg = para.segments[seg_idx]
        before = seg.text[:start]
        middle = seg.text[start : start + length]
        after = seg.text[start + length :]
        replacement: list[_Segment] = []
        if before:
            replacement.append(_Segment(before, seg.tag))
        replacement.append(_Segment(middle, seg.tag))
        if after:
            replacement.append(_Segment(after, seg.tag))
        para.segments[seg_idx : seg_idx + 1] = replacement
        return seg_idx + (1 if before else 0)

    def _apply_replace(self, op: TextReplace) -> None:
        para = self._para(op.node_id)
        self._check_size(para, max(len(op.before), len(op.after)))
        seg_idx, start = self._locate(para, op.before)
        mid = self._split_segment(para, seg_idx, start, len(op.before))
        para.segments[mid].tag = "del"
        para.segments.insert(mid + 1, _Segment(op.after, "ins"))

    def _apply_delete(self, op: TextDelete) -> None:
        para = self._para(op.node_id)
        self._check_size(para, len(op.before))
        seg_idx, start = self._locate(para, op.before)
        mid = self._split_segment(para, seg_idx, start, len(op.before))
        para.segments[mid].tag = "del"

    def _apply_insert(self, op: TextInsert) -> None:
        para = self._para(op.node_id)
        self._check_size(para, len(op.text))
        if op.before_anchor is None:
            para.segments.append(_Segment(op.text, "ins"))
            return
        seg_idx, start = self._locate(para, op.before_anchor)
        # insert AFTER the anchor occurrence
        mid = self._split_segment(para, seg_idx, start, len(op.before_anchor))
        para.segments.insert(mid + 1, _Segment(op.text, "ins"))

    def _apply_style(self, op: StyleChange) -> None:
        para = self._para(op.node_id)
        if op.before != para.style:
            raise BeforeMismatchError(
                f"style before {op.before!r} != current {para.style!r} "
                f"for node {op.node_id!r}"
            )
        para.style = op.style
        para.style_changed = True

    def _apply_node_insert(self, op: NodeInsert) -> None:
        text = str(op.value.get("text", ""))
        style = str(op.value.get("style", "Normal"))
        node_id = self._allocator.mint("p")
        para_id = node_id.split(":", 1)[1].rjust(8, "0").upper()
        new_para = _Para(
            node_id=node_id,
            para_id=para_id,
            style=style,
            base_style=style,
            inserted=True,
            segments=[_Segment(text, "ins")] if text else [],
        )
        pos = max(0, min(op.position, len(self._paras)))
        self._paras.insert(pos, new_para)

    def _apply_node_delete(self, op: NodeDelete) -> None:
        para = self._para(op.node_id)
        para.deleted = True
        for seg in para.segments:
            if seg.tag == "keep":
                seg.tag = "del"

    def _apply_format_run(self, op: FormatRun) -> None:
        """Record a run-property change on a paragraph (guarding ``before``).

        ``before`` must match the paragraph's *current* run props for every key
        the op touches; on a previously-unformatted paragraph the implicit prior
        value is "off"/absent (a falsy/missing entry). The change is stored and
        materialized as a ``w:rPrChange`` at render time.
        """
        para = self._para(op.node_id)
        # before-guard: the op's prior props must match the live props per key.
        for key, prior in op.before.items():
            current = para.format_props.get(key) if para.format_changed else None
            current_norm = bool(current) if current is not None else False
            if bool(prior) != current_norm:
                raise BeforeMismatchError(
                    f"format before {key!r}={prior!r} != current "
                    f"{current_norm!r} for node {op.node_id!r}"
                )
        # Merge: new props take effect; the captured 'before' is the union of the
        # op's declared prior values (so reject restores them).
        merged_before = dict(para.format_before)
        for key in op.props:
            if key not in merged_before:
                merged_before[key] = op.before.get(key, False)
        para.format_props = {**para.format_props, **op.props}
        para.format_before = merged_before
        para.format_changed = True

    def _apply_node_move(self, op: NodeMove) -> None:
        """Move a paragraph via the del+ins surrogate.

        The source paragraph is marked deleted (tracked ``w:del`` runs + a deleted
        pilcrow) and a fresh *inserted* copy carrying the same text/style is placed
        at ``to_index``. Accept-all therefore yields the paragraph at its new spot
        and reject-all yields it at its original spot — a Word-acceptable move
        that reuses proven insert/delete revision machinery.
        """
        source = self._para(op.node_id)
        moved_text = source.baseline_text()
        moved_style = source.style
        # Mark the source as a tracked deletion.
        source.deleted = True
        for seg in source.segments:
            if seg.tag == "keep":
                seg.tag = "del"
        # Materialize the inserted copy at the destination.
        node_id = self._allocator.mint("p")
        para_id = node_id.split(":", 1)[1].rjust(8, "0").upper()
        moved = _Para(
            node_id=node_id,
            para_id=para_id,
            style=moved_style,
            base_style=moved_style,
            inserted=True,
            segments=[_Segment(moved_text, "ins")] if moved_text else [],
        )
        dest = max(0, min(op.to_index, len(self._paras)))
        self._paras.insert(dest, moved)

    # -- render / save --------------------------------------------------------

    def render_tracked(self) -> bytes:
        """Rebuild every paragraph's runs as native revisions and return bytes."""
        import io

        alloc = WidAllocator()
        docx_paras = list(self._doc.paragraphs)
        # Map model paragraphs back onto the live python-docx paragraphs by id,
        # appending new XML paragraphs for node.insert.
        live_by_id = {self._ensure_para_id(p): p for p in docx_paras}
        for para in self._paras:
            target = live_by_id.get(para.para_id)
            if target is None:
                target = self._append_paragraph(para.para_id)
            self._render_paragraph(target, para, alloc)
        buffer = io.BytesIO()
        self._doc.save(buffer)
        return buffer.getvalue()

    def _append_paragraph(self, para_id: str):  # type: ignore[no-untyped-def]
        p = self._doc.add_paragraph()
        p._p.set(qn("w14:paraId"), para_id)
        return p

    def _render_paragraph(self, docx_para, para: _Para, alloc: WidAllocator) -> None:  # type: ignore[no-untyped-def]
        p_el = docx_para._p
        # clear existing runs (keep pPr)
        for run in list(p_el.findall(qn("w:r"))):
            p_el.remove(run)
        if docx_para.style is not None and para.style:
            try:
                docx_para.style = para.style
            except KeyError:
                pass
        ppr = self._ensure_ppr(p_el)
        if para.style_changed:
            ppr.append(
                make_ppr_change(para.base_style, self._author, self._date, alloc)
            )
        if para.inserted:
            ppr.append(make_inserted_pilcrow(self._author, self._date, alloc))
        if para.deleted:
            ppr.append(make_deleted_pilcrow(self._author, self._date, alloc))
        for seg in para.segments:
            if not seg.text:
                continue
            if seg.tag == "ins":
                run = make_inserted_run(seg.text, self._author, self._date, alloc)
            elif seg.tag == "del":
                run = make_deleted_run(seg.text, self._author, self._date, alloc)
            else:
                run = self._plain_run(seg.text)
            if para.format_changed and seg.tag != "del":
                self._apply_run_format(run, para, seg.tag, alloc)
            p_el.append(run)

    def _apply_run_format(self, run, para: _Para, tag: str, alloc: WidAllocator) -> None:  # type: ignore[no-untyped-def]
        """Attach the live run props + a ``w:rPrChange`` to a rendered run.

        ``run`` is either a bare ``w:r`` (keep) or a ``w:ins`` wrapper (ins); the
        actual ``w:r`` to format is the wrapper's child for the inserted case. The
        new props go in the run's ``w:rPr`` (created first so it precedes ``w:t``)
        and the prior props are captured in a nested ``w:rPrChange``.
        """
        from docx.oxml.ns import qn as _qn

        r_el = run if run.tag == _qn("w:r") else run.find(_qn("w:r"))
        if r_el is None:
            return
        rpr = build_run_props(para.format_props)
        rpr.append(make_rpr_change(para.format_before, self._author, self._date, alloc))
        r_el.insert(0, rpr)

    def _ensure_ppr(self, p_el):  # type: ignore[no-untyped-def]
        ppr = p_el.find(qn("w:pPr"))
        if ppr is None:
            ppr = p_el.makeelement(qn("w:pPr"), {})
            p_el.insert(0, ppr)
        return ppr

    def _plain_run(self, text: str):  # type: ignore[no-untyped-def]
        from lxml import etree

        run = etree.Element(qn("w:r"))
        t = etree.SubElement(run, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = text
        return run

    def save(self, out_path: str) -> None:
        """Save the rendered tracked document to ``out_path``."""
        resolved = safe_path(out_path, allow_suffixes=(".docx",))
        resolved.parent.mkdir(parents=True, exist_ok=True)
        resolved.write_bytes(self.render_tracked())

    # -- accessors ------------------------------------------------------------

    @property
    def allocator(self) -> NodeIdAllocator:
        return self._allocator

    def node_id_map(self) -> dict[str, str]:
        """Return the persisted ``node_id -> carrier`` map for the header."""
        return self._allocator.as_map()
