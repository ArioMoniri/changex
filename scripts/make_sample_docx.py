#!/usr/bin/env python3
"""Generate a small sample ``.docx`` fixture for the ChangeX M0 spine.

The fixture is crafted to exercise the linchpin addressing decisions:

* Most paragraphs carry an explicit ``w14:paraId`` so the adapter can prove it
  *reuses* Word's native stable id rather than minting a fresh one.
* Two paragraphs share **identical text** ("This line is intentionally
  duplicated.") to prove node_ids are NOT content hashes — duplicate content
  must still get distinct, stable ids.
* One paragraph deliberately *omits* its paraId so the adapter's mint+inject
  path is covered too.
* A "Heading 1" paragraph is included as the target for a ``style.change`` op.

Run directly to (re)create ``examples/sample.docx``::

    python scripts/make_sample_docx.py [out_path]

The script only depends on ``python-docx`` + ``lxml`` (already required by
``changex-core``); it does not import ``changex_core`` itself, so the fixture is
a neutral input rather than a ChangeX-produced artifact.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Repo-relative default so the example/tests can find the fixture deterministically.
_DEFAULT_OUT = Path(__file__).resolve().parent.parent / "examples" / "sample.docx"


def _set_para_id(paragraph, para_id: str) -> None:
    """Stamp an explicit ``w14:paraId`` onto a python-docx paragraph element."""
    paragraph._p.set(qn("w14:paraId"), para_id)


def build_sample() -> Document:
    """Build the in-memory sample document with the addressing edge cases."""
    doc = Document()

    # A heading paragraph — the target for a style.change op.
    heading = doc.add_paragraph("Quarterly Report", style="Heading 1")
    _set_para_id(heading, "10000001")

    # A body paragraph with edit targets ("quick", "lazy") for text ops.
    body = doc.add_paragraph(
        "The quick brown fox jumps over the lazy dog every single morning."
    )
    _set_para_id(body, "10000002")

    # Two duplicate-content paragraphs: identical text, must get distinct ids.
    dup_a = doc.add_paragraph("This line is intentionally duplicated.")
    _set_para_id(dup_a, "10000003")
    dup_b = doc.add_paragraph("This line is intentionally duplicated.")
    _set_para_id(dup_b, "10000004")

    # A paragraph that DELIBERATELY lacks a paraId (mint+inject path).
    doc.add_paragraph("This paragraph has no native paraId on purpose.")

    # A trailing paragraph used as an insertion anchor / context.
    closing = doc.add_paragraph("Prepared by the analytics team.")
    _set_para_id(closing, "10000006")

    return doc


def write_sample(out_path: Path = _DEFAULT_OUT) -> Path:
    """Build and save the sample document, returning the written path."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_sample()
    doc.save(str(out_path))
    return out_path


def main(argv: list[str] | None = None) -> int:
    args = list(sys.argv[1:] if argv is None else argv)
    out = Path(args[0]) if args else _DEFAULT_OUT
    written = write_sample(out)
    print(f"wrote sample docx: {written}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
