"""External-engine accept/reject oracle (LibreOffice headless).

The strongest possible round-trip check: a *third-party* engine that did not author
the revisions resolves them, and we compare. It complements the always-on portable
resolver in ``test_docx_roundtrip.py``.

Headless tracked-change resolution is only reliable on some platforms — it works on
Linux CI, but the macOS app build no-ops ``macro:///`` and the bundled UNO Python is
killed by platform security. To stay honest and non-flaky this module is:

* marked ``libreoffice`` and **deselected by default** (see ``pytest.ini``), so a
  normal ``pytest`` run never reports it as skipped; CI runs it with ``-m libreoffice``;
* **capability-probed** at runtime — it first runs a trivial macro that writes a
  sentinel file; if the sentinel is not produced, headless macro execution is
  unavailable here and the assertions ``pytest.skip`` rather than fail.

So: on Linux CI it runs for real (genuine external oracle); anywhere it can't, it
skips cleanly; and a real serialization bug still fails it where the engine works.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import textwrap
from pathlib import Path

import pytest

import changex_core as cx
from changex_core.adapters.docx_adapter import DocxAdapter

pytest.importorskip("docx", reason="python-docx is required for the fixture")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

pytestmark = pytest.mark.libreoffice


def _find_soffice() -> str | None:
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    return mac if os.path.exists(mac) else None


_SOFFICE = _find_soffice()


def _install_macro(profile: Path, body: str) -> None:
    """Install a parameterless ``Standard.Module1.Main`` into a throwaway profile."""
    basic = profile / "user" / "basic" / "Standard"
    basic.mkdir(parents=True, exist_ok=True)
    (basic / "Module1.xba").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE script:module PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "module.dtd">\n'
        '<script:module xmlns:script="http://openoffice.org/2000/script" '
        'script:name="Module1" script:language="StarBasic">' + body + "</script:module>",
        encoding="utf-8",
    )
    (basic / "script.xlb").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE library:library PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "library.dtd">\n'
        '<library:library xmlns:library="http://openoffice.org/2000/library" '
        'library:name="Standard" library:readonly="false" library:passwordprotected="false">'
        '<library:element library:name="Module1"/></library:library>',
        encoding="utf-8",
    )
    (profile / "user" / "basic" / "script.xlc").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE library:libraries PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "libraries.dtd">\n'
        '<library:libraries xmlns:library="http://openoffice.org/2000/library" '
        'xmlns:xlink="http://www.w3.org/1999/xlink">'
        '<library:library library:name="Standard" '
        'xlink:href="$(USER)/basic/Standard/script.xlb/" xlink:type="simple" library:link="false"/>'
        "</library:libraries>",
        encoding="utf-8",
    )


def _run(profile: Path) -> None:
    assert _SOFFICE is not None
    subprocess.run(
        [
            _SOFFICE,
            "--headless",
            "--norestore",
            "--nologo",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "macro:///Standard.Module1.Main",
        ],
        capture_output=True,
        timeout=180,
    )


@pytest.fixture(scope="module")
def capable(tmp_path_factory) -> bool:
    """True iff headless macro execution actually works on this machine."""
    if _SOFFICE is None:
        return False
    prof = tmp_path_factory.mktemp("probe")
    sentinel = prof / "probe_ok.txt"
    _install_macro(
        prof,
        textwrap.dedent(
            f"""\
            Sub Main
                Dim i As Integer
                i = FreeFile
                Open "{sentinel.resolve()}" For Output As #i
                Print #i, "ok"
                Close #i
            End Sub
            """
        ),
    )
    _run(prof)
    return sentinel.exists()


def _build_tracked(tmp_path: Path) -> tuple[Path, str, str]:
    base = tmp_path / "base.docx"
    doc = Document()
    p = doc.add_paragraph("The quick brown fox jumps over the lazy dog every single morning.")
    p._p.set(qn("w14:paraId"), "20000001")
    doc.add_paragraph("Prepared by the analytics team.")
    doc.save(str(base))

    a = DocxAdapter.load(str(base))
    m = a.to_model()
    baseline = "\n".join(x.text() for x in m.child_paragraphs())
    body = next(x for x in m.child_paragraphs() if "quick" in x.text())
    a.apply(cx.TextReplace(node_id=body.node_id, before="quick", after="swift"))
    a.apply(cx.TextDelete(node_id=body.node_id, before=" every single morning"))
    active = "\n".join(x.text() for x in a.to_model().child_paragraphs())

    tracked = tmp_path / "tracked.docx"
    a.save(str(tracked))
    return tracked, active, baseline


def _resolve_via_soffice(tracked: Path, out_txt: Path, *, accept: bool, profile: Path) -> str:
    _install_macro(
        profile,
        textwrap.dedent(
            f"""\
            Sub Main
                Dim oArgs(0) As New com.sun.star.beans.PropertyValue
                oArgs(0).Name = "Hidden" : oArgs(0).Value = True
                Dim oDoc As Object
                oDoc = StarDesktop.loadComponentFromURL("{tracked.resolve().as_uri()}", "_blank", 0, oArgs())
                If {"True" if accept else "False"} Then
                    oDoc.AcceptAllTrackedChanges()
                Else
                    oDoc.RejectAllTrackedChanges()
                End If
                Dim oSave(0) As New com.sun.star.beans.PropertyValue
                oSave(0).Name = "FilterName" : oSave(0).Value = "Text"
                oDoc.storeToURL("{out_txt.resolve().as_uri()}", oSave())
                oDoc.close(False)
            End Sub
            """
        ),
    )
    _run(profile)
    if not out_txt.exists():
        pytest.skip("LibreOffice produced no output (headless resolution unavailable here)")
    return out_txt.read_text(encoding="utf-8", errors="replace")


def _norm(s: str) -> str:
    """Collapse whitespace per line so LibreOffice's Text-filter quirks don't matter."""
    return "\n".join(" ".join(line.split()) for line in s.strip().splitlines())


def test_libreoffice_accept_all_equals_active_capture(tmp_path: Path, capable: bool) -> None:
    if not capable:
        pytest.skip("LibreOffice present but headless macro execution unavailable (e.g. macOS)")
    tracked, active, _baseline = _build_tracked(tmp_path)
    got = _resolve_via_soffice(tracked, tmp_path / "accepted.txt", accept=True, profile=tmp_path / "pa")
    assert _norm(got) == _norm(active)


def test_libreoffice_reject_all_equals_baseline(tmp_path: Path, capable: bool) -> None:
    if not capable:
        pytest.skip("LibreOffice present but headless macro execution unavailable (e.g. macOS)")
    tracked, _active, baseline = _build_tracked(tmp_path)
    got = _resolve_via_soffice(tracked, tmp_path / "rejected.txt", accept=False, profile=tmp_path / "pr")
    assert _norm(got) == _norm(baseline)
