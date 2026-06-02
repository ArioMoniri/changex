"""Correctness follow-ups: journal-aware save + baseline-bound verify.

(a) ``save_active`` must save a tracked .docx that is a pure projection of the
    journal's NON-reverted events, so a reverted op's revision is genuinely gone
    from the saved file (not merely flagged in the journal).

(b) ``Journal.verify`` must re-hash the on-disk baseline against the header's
    ``baseline_sha256`` and set ``VerifyResult.baseline_match`` accordingly.
"""

from __future__ import annotations

from pathlib import Path

import pytest

import changex_core as cx
from changex_core.adapters.docx_adapter import DocxAdapter
from changex_core.journal.events import Header, Provenance, Target, utc_now_iso
from changex_core.render.save import save_active

docx = pytest.importorskip("docx", reason="python-docx required for save/verify tests")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


def _baseline(path: Path) -> str:
    doc = Document()
    p = doc.add_paragraph("The quick brown fox jumps.")
    p._p.set(qn("w14:paraId"), "10000001")
    doc.save(str(path))
    return "p:10000001"


def _all_text_fragments(path: Path) -> tuple[str, list[str]]:
    """Return (joined text, per-element [(tag, text)]) over w:t/w:delText."""
    doc = Document(str(path))
    frags: list[str] = []
    joined: list[str] = []
    for para in doc.paragraphs:
        for el in para._p.iter():
            if el.tag in (qn("w:t"), qn("w:delText")) and el.text:
                tag = el.tag.split("}")[-1]
                frags.append(f"{tag}:{el.text}")
                joined.append(el.text)
    return " ".join(joined), frags


def _open_with_two_ops(tmp_path: Path):  # type: ignore[no-untyped-def]
    base = tmp_path / "doc.docx"
    node_id = _baseline(base)
    adapter = DocxAdapter.load(str(base))
    header = Header.create(
        baseline_sha256=adapter.baseline_sha256(),
        filename=base.name,
        baseline_uri=str(base),
        node_id_map=adapter.node_id_map(),
    )
    journal = cx.Journal.open(str(tmp_path / "s.changex"), header=header)
    sid = journal.header.session_id

    def prov() -> Provenance:
        return Provenance(
            ts=utc_now_iso(), session_id=sid, agent="m", vendor="cli", provenance_source="declared"
        )

    def tgt() -> Target:
        return Target(node_id=node_id, node_kind="paragraph", path="")

    op1 = cx.TextReplace(node_id=node_id, before="quick", after="swift")
    op2 = cx.TextReplace(node_id=node_id, before="brown", after="crimson")
    adapter.apply(op1)
    e1 = journal.append(op1, tgt(), prov())
    adapter.apply(op2)
    e2 = journal.append(op2, tgt(), prov())
    return base, journal, e1, e2


# --------------------------------------------------------------------------- #
# (a) Journal-aware save drops a reverted op's revision from the saved docx.
# --------------------------------------------------------------------------- #
def test_save_active_includes_all_ops_when_none_reverted(tmp_path: Path) -> None:
    base, journal, _e1, _e2 = _open_with_two_ops(tmp_path)
    out = tmp_path / "out.docx"
    n = save_active(journal, str(base), str(out))
    assert n == 2
    joined, _ = _all_text_fragments(out)
    assert "swift" in joined and "crimson" in joined


def test_save_active_drops_reverted_op_revision_from_saved_docx(tmp_path: Path) -> None:
    base, journal, _e1, e2 = _open_with_two_ops(tmp_path)
    journal.revert(e2.op_id)  # reject brown -> crimson

    out = tmp_path / "out.docx"
    n = save_active(journal, str(base), str(out))
    assert n == 1  # only the active op replayed

    joined, _frags = _all_text_fragments(out)
    # The reverted insertion must NOT be present anywhere in the saved file...
    assert "crimson" not in joined
    # ...the active insertion IS present, and 'brown' is restored as plain text.
    assert "swift" in joined
    assert "brown" in joined


# --------------------------------------------------------------------------- #
# (b) verify() re-hashes the on-disk baseline to set baseline_match.
# --------------------------------------------------------------------------- #
def test_verify_baseline_match_true_for_unmodified_baseline(tmp_path: Path) -> None:
    base, journal, _e1, _e2 = _open_with_two_ops(tmp_path)
    result = journal.verify(baseline_path=str(base))
    assert result.ok is True
    assert result.baseline_match is True


def test_verify_baseline_match_uses_header_uri_when_no_arg(tmp_path: Path) -> None:
    base, journal, _e1, _e2 = _open_with_two_ops(tmp_path)
    # header.baseline_uri points at `base`; verify with no arg should use it.
    result = journal.verify()
    assert result.baseline_match is True
    assert "baseline matches" in result.detail


def test_verify_baseline_match_false_when_baseline_drifts(tmp_path: Path) -> None:
    base, journal, _e1, _e2 = _open_with_two_ops(tmp_path)
    # Mutate the on-disk baseline so its sha256 no longer matches the header.
    drifted = Document(str(base))
    drifted.add_paragraph("An out-of-band edit to the baseline file.")
    drifted.save(str(base))

    result = journal.verify(baseline_path=str(base))
    # Chain is still intact (ok), but the baseline binding is broken.
    assert result.ok is True
    assert result.baseline_match is False
    assert "baseline mismatch" in result.detail


def test_verify_baseline_unchecked_when_no_baseline_available(tmp_path: Path) -> None:
    # A journal whose header has a baseline_sha256 but no readable baseline.
    header = Header.create(baseline_sha256="0" * 64, filename="ghost.docx")
    journal = cx.Journal.open(str(tmp_path / "ghost.changex"), header=header)
    result = journal.verify()
    # Unknown baseline is not a failure; match stays True with a noting detail.
    assert result.ok is True
    assert result.baseline_match is True
    assert "not checked" in result.detail
