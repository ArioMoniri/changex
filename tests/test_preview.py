"""Tests for ``changex preview`` — cross-platform HTML preview (journal redline + code)."""

from __future__ import annotations

import json
from pathlib import Path

import pytest

from changex_core.preview import preview_html


def _real_journal(tmp_path: Path) -> Path:
    """Build a schema-valid, hash-chained .changex via the real `track` flow."""
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    base = tmp_path / "base.docx"
    doc = docx.Document()
    p = doc.add_paragraph("The quick brown fox.")
    p._p.set(qn("w14:paraId"), "20000001")
    doc.save(str(base))

    import changex_core as cx
    from changex_core.cli import main

    adapter = cx.DocxAdapter.load(str(base))
    node = next(x for x in adapter.to_model().child_paragraphs() if "quick" in x.text())
    ops = [{"kind": "text.replace", "node_id": node.node_id, "before": "quick", "after": "swift"}]
    ops_path = tmp_path / "ops.json"
    ops_path.write_text(json.dumps(ops), encoding="utf-8")

    journal = tmp_path / "x.changex"
    tracked = tmp_path / "tracked.docx"
    rc = main(["track", str(base), str(ops_path), "--out", str(tracked), "--changex", str(journal)])
    assert rc == 0
    return journal


def test_preview_changex_renders_redline(tmp_path: Path) -> None:
    html = preview_html(_real_journal(tmp_path))
    assert "<!doctype html" in html.lower()
    assert "swift" in html and "quick" in html  # insertion + deletion
    assert "<ins" in html or "<del" in html


def test_preview_code_is_solid_and_contains_source(tmp_path: Path) -> None:
    f = tmp_path / "sample.py"
    f.write_text("import math\n\ndef f(x):\n    return x * 2\n", encoding="utf-8")
    html = preview_html(f)
    assert "background:#ffffff" in html  # solid background — never blank
    assert "def" in html and "math" in html


def test_preview_code_is_highlighted_when_pygments_present(tmp_path: Path) -> None:
    pytest.importorskip("pygments")
    f = tmp_path / "sample.py"
    f.write_text("import math\n", encoding="utf-8")
    html = preview_html(f)
    # Pygments emits a .highlight container and token <span> classes.
    assert 'class="highlight"' in html
    assert "<span" in html


def test_preview_plain_fallback_without_pygments(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import builtins

    real_import = builtins.__import__

    def no_pygments(name: str, *args: object, **kwargs: object) -> object:
        if name == "pygments" or name.startswith("pygments."):
            raise ImportError("pygments disabled for test")
        return real_import(name, *args, **kwargs)  # type: ignore[arg-type]

    monkeypatch.setattr(builtins, "__import__", no_pygments)
    f = tmp_path / "x.txt"
    f.write_text("plain <text> & more\n", encoding="utf-8")
    html = preview_html(f)
    assert "<pre>" in html
    assert "&lt;text&gt;" in html  # escaped, not raw


def test_preview_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(Exception):
        preview_html(tmp_path / "nope.py")


def test_log_and_kraken_graph_render(tmp_path: Path) -> None:
    """changex log (git-style) and the GitKraken-style commit graph render from a journal."""
    from changex_core.journal.journal import Journal
    from changex_core.render.html import render_html, render_log

    journal_path = _real_journal(tmp_path)
    j = Journal.open(str(journal_path))
    events = j.active_events()

    log = render_log(events, oneline=True)
    assert "text.replace" in log and "quick" in log and "swift" in log
    full = render_log(events, oneline=False)
    assert "commit " in full and "Author:" in full and "Date:" in full

    graph = render_html(events, header=j.header.to_dict())
    assert "kx-head" in graph and 'class="commit"' in graph  # the commit graph
    assert "<ins>" in graph and "<del>" in graph              # the redline
    assert ".rail" in graph and ".node" in graph              # the graph rail + dots


def test_rewind_reconstructs_earlier_phases(tmp_path: Path) -> None:
    """`changex rewind --to N` rebuilds the document at an earlier phase of its history."""
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    from changex_core.cli import main

    journal = _real_journal(tmp_path)          # one edit: replace "quick" → "swift"
    base = tmp_path / "base.docx"

    def revisions(path: Path) -> int:
        d = docx.Document(str(path))
        return sum(
            1
            for para in d.paragraphs
            for el in para._p.iter()
            if el.tag in (qn("w:ins"), qn("w:del"))
        )

    out0 = tmp_path / "rw0.docx"
    assert main(["rewind", str(journal), str(base), "--to", "0", "--out", str(out0)]) == 0
    assert revisions(out0) == 0                # phase 0 = clean baseline

    out1 = tmp_path / "rw1.docx"
    assert main(["rewind", str(journal), str(base), "--to", "1", "--out", str(out1)]) == 0
    assert revisions(out1) > 0                 # phase 1 carries the edit
