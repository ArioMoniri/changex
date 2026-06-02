"""Passive (``open`` / ``seal``) capture tests: reconstruct a hand-edited docx.

These exercise the "native to any model" path end-to-end: snapshot a baseline,
edit the docx with plain python-docx (standing in for an arbitrary model / tool /
human that does NO tool-calling), seal, and assert the journal now carries a
best-effort op stream with **honestly degraded** provenance.

The diff layer is also unit-tested directly so each op kind's reconstruction is
pinned, independent of docx parsing.
"""

from __future__ import annotations

from pathlib import Path

import pytest

import changex_core as cx
from changex_core.adapters.docx_adapter import DocxAdapter
from changex_core.diff.text_diff import ParagraphSpec, diff_paragraphs
from changex_core.passive import (
    PASSIVE_RATIONALE,
    open_passive,
    seal_passive,
)

docx = pytest.importorskip("docx", reason="python-docx is required for passive docx tests")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


# --------------------------------------------------------------------------- #
# Diff layer: each op kind reconstructs from isolated paragraph edits.
# --------------------------------------------------------------------------- #
def _spec(node_id, text, style="Normal"):  # type: ignore[no-untyped-def]
    return ParagraphSpec(node_id=node_id, text=text, style=style)


def test_diff_reconstructs_text_replace() -> None:
    diff = diff_paragraphs([_spec("p:1", "the quick fox")], [_spec("p:1", "the swift fox")])
    assert [o.op["kind"] for o in diff.ops] == ["text.replace"]
    op = diff.ops[0].op
    assert op["before"] == "quick" and op["after"] == "swift"
    assert op["node_id"] == "p:1"
    assert diff.replaced == 1


def test_diff_reconstructs_text_delete_and_insert() -> None:
    deleted = diff_paragraphs([_spec("p:1", "hello brave world")], [_spec("p:1", "hello world")])
    assert [o.op["kind"] for o in deleted.ops] == ["text.delete"]
    assert deleted.ops[0].op["before"] == "brave "

    inserted = diff_paragraphs([_spec("p:1", "hello world")], [_spec("p:1", "hello brave world")])
    assert [o.op["kind"] for o in inserted.ops] == ["text.insert"]
    assert inserted.ops[0].op["text"] == "brave "


def test_diff_reconstructs_node_delete_and_insert() -> None:
    base = [_spec("p:1", "alpha"), _spec("p:2", "beta"), _spec("p:3", "gamma")]
    removed = diff_paragraphs(base, [_spec("p:1", "alpha"), _spec("p:3", "gamma")])
    assert [o.op["kind"] for o in removed.ops] == ["node.delete"]
    assert removed.ops[0].op["node_id"] == "p:2"
    assert removed.ops[0].op["value"]["text"] == "beta"
    assert removed.deleted == 1

    added = diff_paragraphs(
        [_spec("p:1", "alpha")],
        [_spec("p:1", "alpha"), _spec(None, "brand new tail")],
    )
    assert [o.op["kind"] for o in added.ops] == ["node.insert"]
    assert added.ops[0].op["position"] == 1
    assert added.ops[0].op["value"]["text"] == "brand new tail"
    assert added.inserted == 1


def test_diff_reconstructs_style_change() -> None:
    diff = diff_paragraphs(
        [_spec("p:1", "Title", "Normal")], [_spec("p:1", "Title", "Heading 1")]
    )
    assert [o.op["kind"] for o in diff.ops] == ["style.change"]
    assert diff.ops[0].op == {
        "kind": "style.change",
        "node_id": "p:1",
        "style": "Heading 1",
        "before": "Normal",
    }


def test_diff_noop_on_identical_sequences() -> None:
    diff = diff_paragraphs([_spec("p:1", "same")], [_spec("p:1", "same")])
    assert diff.ops == []
    assert diff.total == 0


# --------------------------------------------------------------------------- #
# End-to-end: open -> hand-edit the same docx in place -> seal.
# --------------------------------------------------------------------------- #
def _write_baseline(path: Path) -> None:
    doc = Document()
    p0 = doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    p0._p.set(qn("w14:paraId"), "10000001")
    p1 = doc.add_paragraph("This second paragraph is unchanged.")
    p1._p.set(qn("w14:paraId"), "10000002")
    doc.save(str(path))


def test_open_writes_passive_header_and_baseline_sidecar(tmp_path: Path) -> None:
    base = tmp_path / "doc.docx"
    _write_baseline(base)

    result = open_passive(str(base))
    assert result.changex_path.exists()
    assert result.paragraphs == 2

    journal = cx.Journal.open(str(result.changex_path))
    assert journal.header.session.get("capture_mode") == "passive"
    # The baseline bytes are preserved out-of-band so seal can diff in place.
    sidecar = Path(journal.header.doc["baseline_uri"])
    assert sidecar.exists()
    assert journal.header.baseline_sha256 == result.baseline.sha256
    # No ops yet — open is purely a snapshot + header.
    assert journal.active_events() == []


def test_seal_reconstructs_hand_edited_docx_with_degraded_provenance(tmp_path: Path) -> None:
    base = tmp_path / "doc.docx"
    _write_baseline(base)
    result = open_passive(str(base))

    # Edit the SAME file in place — exactly what a no-tool-calling model/tool does.
    edited = Document(str(base))
    for run in edited.paragraphs[0].runs:
        run.text = run.text.replace("quick", "swift")
    edited.save(str(base))

    seal = seal_passive(str(base))
    assert seal.degraded is True
    assert seal.baseline_unchanged is False
    assert seal.appended >= 1
    assert seal.replaced >= 1

    journal = cx.Journal.open(str(result.changex_path))
    events = journal.active_events()
    assert events, "seal should have appended reconstructed ops"

    replace_events = [e for e in events if e.op.get("kind") == "text.replace"]
    assert replace_events, "expected a reconstructed text.replace"
    op = replace_events[0].op
    assert op["before"] == "quick" and op["after"] == "swift"

    # Provenance must be honestly degraded on every reconstructed op.
    for event in events:
        prov = event.provenance
        assert prov.provenance_source == "observed"
        assert prov.rationale == PASSIVE_RATIONALE
        assert prov.agent is None
        assert prov.vendor is None
        assert prov.turn_id is None
        assert prov.prompt_sha256 is None


def test_sealed_ops_replay_cleanly_onto_baseline(tmp_path: Path) -> None:
    base = tmp_path / "doc.docx"
    _write_baseline(base)
    result = open_passive(str(base))

    edited = Document(str(base))
    for run in edited.paragraphs[0].runs:
        run.text = run.text.replace("quick", "swift")
    edited.save(str(base))
    seal_passive(str(base))

    journal = cx.Journal.open(str(result.changex_path))
    baseline_uri = journal.header.doc["baseline_uri"]
    adapter = DocxAdapter.load(baseline_uri)
    baseline_model = adapter.to_model()
    replayed = journal.replay(adapter, baseline_model)
    texts = [n.text() for n in replayed.children if n.node_kind == cx.NodeKind.PARAGRAPH]
    assert texts[0] == "The swift brown fox jumps over the lazy dog."
    assert texts[1] == "This second paragraph is unchanged."


def test_seal_noop_when_unedited(tmp_path: Path) -> None:
    base = tmp_path / "doc.docx"
    _write_baseline(base)
    open_passive(str(base))
    seal = seal_passive(str(base))
    assert seal.baseline_unchanged is True
    assert seal.appended == 0


def test_seal_rejects_active_capture_journal(tmp_path: Path) -> None:
    """seal must refuse a journal that was not opened in passive mode."""
    base = tmp_path / "doc.docx"
    _write_baseline(base)
    adapter = DocxAdapter.load(str(base))
    header = cx.Header.create(
        baseline_sha256=adapter.baseline_sha256(),
        filename=base.name,
        capture_mode="active",
    )
    changex = tmp_path / "active.changex"
    cx.Journal.open(str(changex), header=header)
    with pytest.raises(ValueError):
        seal_passive(str(base), str(changex))
