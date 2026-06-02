"""LibreOffice (soffice) native accept/reject round-trip — GATED, skips if absent.

This is the *external-tool* counterpart to the in-process accept/reject projection
already covered by ``test_docx_adapter.py``. It proves the ``w:ins`` / ``w:del``
revisions ChangeX writes are **real Word revisions a third-party engine resolves
the same way ChangeX projects them**:

* **accept-all** (in LibreOffice) == the adapter's active-capture target text, and
* **reject-all** (in LibreOffice) == the baseline text.

It is deliberately *gating*: if neither ``soffice`` nor ``libreoffice`` is on the
PATH the whole module is skipped (``pytest.skip``) — we never silently pass a
round-trip we did not actually run, and we never fail CI for a missing optional
tool. When LibreOffice *is* present, the two equalities above must hold.

Mechanism
---------
LibreOffice has no stable CLI flag to accept/reject tracked changes, so we drive
it headlessly with a tiny Basic macro (``--headless ... macro://``) that opens
the tracked ``.docx``, calls ``AcceptAllTrackedChanges`` /
``RejectAllTrackedChanges`` on the document, and saves a flat ``.txt`` we can
read back and compare. The macro is installed into a throwaway user profile so
the host's real LibreOffice config is never touched.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import textwrap
from pathlib import Path

import pytest

import changex_core as cx
from changex_core.adapters.docx_adapter import DocxAdapter

# --------------------------------------------------------------------------- #
# Tool discovery / gating
# --------------------------------------------------------------------------- #
pytest.importorskip("docx", reason="python-docx is required to build the fixture")


def _find_soffice() -> str | None:
    """Return a usable ``soffice``/``libreoffice`` executable path, or None.

    Checks the PATH first (``soffice`` then ``libreoffice``) and then the
    well-known macOS app-bundle location so the test still runs on a typical
    Mac install that does not symlink ``soffice`` onto the PATH.
    """
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    mac_bundle = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.exists(mac_bundle):
        return mac_bundle
    return None


_SOFFICE = _find_soffice()

pytestmark = pytest.mark.skipif(
    _SOFFICE is None,
    reason="LibreOffice (soffice/libreoffice) not on PATH — round-trip skipped",
)


# --------------------------------------------------------------------------- #
# Fixture: a tracked .docx with REAL w:ins / w:del revisions
# --------------------------------------------------------------------------- #
def _build_tracked_docx(tmp_path: Path) -> tuple[Path, str, str]:
    """Build a tracked ``.docx`` and return ``(path, active_text, baseline_text)``.

    The two expected texts are taken from the adapter *before* save: the
    active-capture target is the current model text (accept-all), and the
    baseline is the model text before any ops were applied (reject-all). Both are
    newline-joined paragraph texts so they compare cleanly against LibreOffice's
    flat-text export.
    """
    from docx import Document
    from docx.oxml.ns import qn

    base = tmp_path / "base.docx"
    doc = Document()
    p = doc.add_paragraph(
        "The quick brown fox jumps over the lazy dog every single morning."
    )
    p._p.set(qn("w14:paraId"), "20000001")
    doc.add_paragraph("Prepared by the analytics team.")
    doc.save(str(base))

    adapter = DocxAdapter.load(str(base))
    model = adapter.to_model()
    baseline_text = "\n".join(p.text() for p in model.child_paragraphs())

    body = next(
        p for p in model.child_paragraphs() if "quick brown fox" in p.text()
    )
    adapter.apply(cx.TextReplace(node_id=body.node_id, before="quick", after="swift"))
    adapter.apply(cx.TextDelete(node_id=body.node_id, before=" every single morning"))

    active_text = "\n".join(p.text() for p in adapter.to_model().child_paragraphs())

    tracked = tmp_path / "tracked.docx"
    adapter.save(str(tracked))

    # Sanity: the saved file really carries native revisions (else the round-trip
    # would be meaningless — there would be nothing for soffice to accept/reject).
    xml = Document(str(tracked)).element.xml
    assert "<w:ins" in xml and "<w:del" in xml, "tracked docx lacks w:ins/w:del"
    return tracked, active_text, baseline_text


# --------------------------------------------------------------------------- #
# LibreOffice headless driver
# --------------------------------------------------------------------------- #
_MACRO = textwrap.dedent(
    """\
    Sub ResolveAndExport(sUrl As String, sOutUrl As String, bAccept As Boolean)
        Dim oDoc As Object
        Dim oArgs(0) As New com.sun.star.beans.PropertyValue
        oArgs(0).Name = "Hidden"
        oArgs(0).Value = True
        oDoc = StarDesktop.loadComponentFromURL(sUrl, "_blank", 0, oArgs())
        If bAccept Then
            oDoc.AcceptAllTrackedChanges()
        Else
            oDoc.RejectAllTrackedChanges()
        End If
        Dim oSave(0) As New com.sun.star.beans.PropertyValue
        oSave(0).Name = "FilterName"
        oSave(0).Value = "Text"
        oDoc.storeToURL(sOutUrl, oSave())
        oDoc.close(False)
    End Sub
    """
)


def _path_to_url(path: Path) -> str:
    return Path(path).resolve().as_uri()


def _run_soffice_resolve(
    soffice: str, tracked: Path, out_txt: Path, *, accept: bool, profile: Path
) -> str:
    """Headlessly accept/reject all revisions in ``tracked`` -> flat text string.

    Installs the Basic macro into a throwaway ``profile`` user dir, invokes it via
    ``macro://`` over the tracked doc, and returns the exported plain text with
    trailing blank lines stripped (LibreOffice's Text filter can append a final
    newline).
    """
    # Install the macro into the user profile's Standard Basic library.
    basic_dir = profile / "user" / "basic" / "Standard"
    basic_dir.mkdir(parents=True, exist_ok=True)
    (basic_dir / "Module1.xba").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE script:module PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "module.dtd">\n'
        '<script:module xmlns:script="http://openoffice.org/2000/script" '
        'script:name="Module1" script:language="StarBasic">'
        + _MACRO
        + "</script:module>",
        encoding="utf-8",
    )
    (basic_dir / "script.xlb").write_text(
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<!DOCTYPE library:library PUBLIC "-//OpenOffice.org//DTD OfficeDocument 1.0//EN" "library.dtd">\n'
        '<library:library xmlns:library="http://openoffice.org/2000/library" '
        'library:name="Standard" library:readonly="false" library:passwordprotected="false">'
        '<library:element library:name="Module1"/></library:library>',
        encoding="utf-8",
    )

    macro = (
        "macro:///Standard.Module1.ResolveAndExport("
        f'"{_path_to_url(tracked)}","{_path_to_url(out_txt)}",'
        f"{'True' if accept else 'False'})"
    )
    cmd = [
        soffice,
        "--headless",
        "--norestore",
        "--nologo",
        f"-env:UserInstallation={_path_to_url(profile)}",
        macro,
    ]
    subprocess.run(cmd, check=True, capture_output=True, timeout=120)
    text = out_txt.read_text(encoding="utf-8", errors="replace")
    return "\n".join(line.rstrip("\r") for line in text.splitlines()).strip("\n")


# --------------------------------------------------------------------------- #
# The gating round-trip test
# --------------------------------------------------------------------------- #
def test_libreoffice_accept_all_equals_active_capture(tmp_path: Path) -> None:
    """soffice accept-all of the tracked docx == the adapter's active target."""
    assert _SOFFICE is not None  # guarded by pytestmark
    tracked, active_text, _baseline = _build_tracked_docx(tmp_path)
    out = tmp_path / "accepted.txt"
    got = _run_soffice_resolve(
        _SOFFICE, tracked, out, accept=True, profile=tmp_path / "prof_accept"
    )
    assert got == active_text


def test_libreoffice_reject_all_equals_baseline(tmp_path: Path) -> None:
    """soffice reject-all of the tracked docx == the baseline text."""
    assert _SOFFICE is not None  # guarded by pytestmark
    tracked, _active, baseline_text = _build_tracked_docx(tmp_path)
    out = tmp_path / "rejected.txt"
    got = _run_soffice_resolve(
        _SOFFICE, tracked, out, accept=False, profile=tmp_path / "prof_reject"
    )
    assert got == baseline_text
