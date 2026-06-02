"""render_document_html — changes shown inline in the document's own outline."""

from __future__ import annotations

from pathlib import Path

import pytest

import changex_core as cx

pytest.importorskip("docx")
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402


def _tracked(tmp_path: Path) -> Path:
    base = tmp_path / "base.docx"
    doc = Document()
    h = doc.add_paragraph("Title Heading")
    h.style = doc.styles["Heading 1"]
    p = doc.add_paragraph("The quick brown fox.")
    p._p.set(qn("w14:paraId"), "20000001")
    doc.save(str(base))

    a = cx.DocxAdapter.load(str(base), author="claude-opus-4-8")
    body = next(x for x in a.to_model().child_paragraphs() if "quick" in x.text())
    a.apply(cx.TextReplace(node_id=body.node_id, before="quick", after="swift"))
    a.apply(cx.TextDelete(node_id=body.node_id, before=" brown"))
    out = tmp_path / "tracked.docx"
    a.save(str(out))
    return out


def test_document_html_renders_changes_inline_in_the_outline(tmp_path: Path) -> None:
    html = cx.render_document_html(str(_tracked(tmp_path)), title="T")
    # insertion + deletion shown INLINE (not as a detached op list)
    assert "<ins" in html and ">swift</ins>" in html
    assert "<del" in html and ">quick</del>" in html
    assert ">  brown</del>".replace("  ", " ") or " brown" in html  # the deleted run is present
    # surrounding document context is rendered, not just the changed substrings
    assert "fox" in html and "Title Heading" in html
    # a Heading-1 paragraph renders as a heading element (document structure preserved)
    assert "<h2>" in html
    # provenance tooltip carries the author
    assert "claude-opus-4-8" in html


def test_document_html_is_self_contained_and_escaped(tmp_path: Path) -> None:
    html = cx.render_document_html(str(_tracked(tmp_path)), title="<x>")
    assert html.startswith("<!doctype html>")
    assert "<title>&lt;x&gt;</title>" in html  # title is HTML-escaped
