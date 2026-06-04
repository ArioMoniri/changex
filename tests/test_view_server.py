"""Review-server smoke tests: 200 + redline markup, live accept/reject.

The server is started on an **ephemeral** port (``port=0``), driven over real
HTTP with stdlib :mod:`urllib`, and shut down — so the test exercises the actual
:mod:`http.server` request path, not just the handler in isolation. The server
binds to ``127.0.0.1`` only (asserted), never a routable interface.
"""

from __future__ import annotations

import json
import threading
import urllib.request
from pathlib import Path
from typing import Any

import pytest

import changex_core as cx
from changex_core.journal.events import Header, Provenance, Target, utc_now_iso
from changex_core.render.server import HOST, build_server


def _journal_with_ops(path: Path) -> tuple[cx.Journal, Any, Any]:
    header = Header.create(baseline_sha256="0" * 64, filename="f.docx")
    journal = cx.Journal.open(str(path), header=header)
    sid = journal.header.session_id

    def prov(agent: str) -> Provenance:
        return Provenance(
            ts=utc_now_iso(),
            session_id=sid,
            agent=agent,
            vendor="cli",
            provenance_source="declared",
        )

    def tgt() -> Target:
        return Target(node_id="p:1", node_kind="paragraph", path="")

    e1 = journal.append(
        cx.TextReplace(node_id="p:1", before="quick", after="swift"), tgt(), prov("claude")
    )
    e2 = journal.append(
        cx.TextInsert(node_id="p:1", before_anchor="fox", text=" today"), tgt(), prov("gpt")
    )
    return journal, e1, e2


def _get(url: str) -> tuple[int, bytes]:
    with urllib.request.urlopen(url) as resp:  # noqa: S310 - localhost only
        return resp.status, resp.read()


def _post_json(url: str, payload: dict[str, Any]) -> tuple[int, dict[str, Any]]:
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req) as resp:  # noqa: S310 - localhost only
        return resp.status, json.loads(resp.read())


def test_server_binds_localhost_only(tmp_path: Path) -> None:
    journal, _, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0)
    try:
        host, _port = server.server_address[0], server.server_address[1]
        assert host == HOST == "127.0.0.1"
    finally:
        server.server_close()


def test_get_root_returns_200_and_review_page(tmp_path: Path) -> None:
    journal, _, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0)
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        status, body = _get(f"http://{HOST}:{port}/")
        assert status == 200
        text = body.decode("utf-8")
        assert "<!doctype html>" in text.lower()
        assert 'id="redline"' in text  # the live redline mount point
        assert 'id="timeline"' in text  # the provenance timeline mount point
    finally:
        server.shutdown()
        server.server_close()


def test_redline_endpoint_returns_redline_markup(tmp_path: Path) -> None:
    journal, _, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0)
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        status, body = _get(f"http://{HOST}:{port}/api/redline")
        assert status == 200
        html = json.loads(body)["html"]
        # The redline reuses render_html over active_events -> <ins>/<del> markup.
        assert "<ins>" in html
        assert "swift" in html
    finally:
        server.shutdown()
        server.server_close()


def test_events_endpoint_lists_models_and_counts(tmp_path: Path) -> None:
    journal, _, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0)
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        status, body = _get(f"http://{HOST}:{port}/api/events")
        assert status == 200
        data = json.loads(body)
        assert data["total_count"] == 2
        assert data["active_count"] == 2
        assert set(data["models"]) == {"claude", "gpt"}
    finally:
        server.shutdown()
        server.server_close()


def test_reject_then_accept_re_renders_redline_live(tmp_path: Path) -> None:
    journal, e1, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0)
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base = f"http://{HOST}:{port}"
    try:
        # Reject e1 (quick -> swift): redline must drop 'swift'.
        status, data = _post_json(base + "/api/review", {"op_id": e1.op_id, "action": "reject"})
        assert status == 200
        assert data["active_count"] == 1
        assert "swift" not in data["redline_html"]
        assert journal.is_reverted(e1.op_id)

        # Accept e1 (un-revert): redline brings 'swift' back.
        status, data = _post_json(base + "/api/review", {"op_id": e1.op_id, "action": "accept"})
        assert status == 200
        assert data["active_count"] == 2
        assert "swift" in data["redline_html"]
        assert not journal.is_reverted(e1.op_id)
    finally:
        server.shutdown()
        server.server_close()


def _docx_journal_with_ops(tmp_path: Path) -> tuple[cx.Journal, Path, Any, Any]:
    """Build a baseline-bound .docx journal (header carries baseline_uri).

    Returns (journal, baseline_docx_path, e1, e2). The baseline_uri in the header
    is what the server replays onto to regenerate the in-document outline.
    """
    pytest.importorskip("docx", reason="python-docx required for the in-document view")
    from docx import Document
    from docx.oxml.ns import qn

    from changex_core.adapters.docx_adapter import DocxAdapter

    base = tmp_path / "doc.docx"
    doc = Document()
    h = doc.add_paragraph("Title Heading")
    h.style = doc.styles["Heading 1"]
    p = doc.add_paragraph("The quick brown fox.")
    p._p.set(qn("w14:paraId"), "10000001")
    doc.save(str(base))
    node_id = "p:10000001"

    adapter = DocxAdapter.load(str(base), author="claude-opus-4-8")
    header = Header.create(
        baseline_sha256=adapter.baseline_sha256(),
        filename=base.name,
        baseline_uri=str(base),
        node_id_map=adapter.node_id_map(),
    )
    journal = cx.Journal.open(str(tmp_path / "s.changex"), header=header)
    sid = journal.header.session_id

    def prov(agent: str) -> Provenance:
        return Provenance(
            ts=utc_now_iso(), session_id=sid, agent=agent, vendor="cli",
            provenance_source="declared",
        )

    def tgt() -> Target:
        return Target(node_id=node_id, node_kind="paragraph", path="")

    op1 = cx.TextReplace(node_id=node_id, before="quick", after="swift")
    op2 = cx.TextDelete(node_id=node_id, before=" brown")
    adapter.apply(op1)
    e1 = journal.append(op1, tgt(), prov("claude"))
    adapter.apply(op2)
    e2 = journal.append(op2, tgt(), prov("gpt"))
    return journal, base, e1, e2


def test_docx_doc_path_renders_in_document_outline(tmp_path: Path) -> None:
    journal, base, _e1, _e2 = _docx_journal_with_ops(tmp_path)
    server = build_server(journal, port=0, doc_path=str(base))
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        # The page heading reflects the in-document outline mode.
        status, body = _get(f"http://{HOST}:{port}/")
        assert status == 200
        assert "Document outline" in body.decode("utf-8")

        # The redline pane shows the full document outline (heading + surrounding
        # context), with the change inline -- not the op-by-op log.
        status, body = _get(f"http://{HOST}:{port}/api/redline")
        assert status == 200
        html = json.loads(body)["html"]
        assert "Title Heading" in html  # document context, not just the change
        assert "<h2>" in html  # Heading-1 paragraph -> outline heading
        assert ">swift</ins>" in html  # insertion shown inline
        assert "fox" in html  # surrounding paragraph text rendered
    finally:
        server.shutdown()
        server.server_close()


def test_docx_in_document_outline_re_renders_on_reject(tmp_path: Path) -> None:
    journal, base, e1, _e2 = _docx_journal_with_ops(tmp_path)
    server = build_server(journal, port=0, doc_path=str(base))
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base_url = f"http://{HOST}:{port}"
    try:
        # Reject e1 (quick -> swift): regenerated outline must drop 'swift'.
        status, data = _post_json(base_url + "/api/review", {"op_id": e1.op_id, "action": "reject"})
        assert status == 200
        assert "Title Heading" in data["redline_html"]  # still the in-document outline
        assert "swift" not in data["redline_html"]
        assert journal.is_reverted(e1.op_id)

        # Accept e1: 'swift' comes back inline in the regenerated outline.
        status, data = _post_json(base_url + "/api/review", {"op_id": e1.op_id, "action": "accept"})
        assert status == 200
        assert ">swift</ins>" in data["redline_html"]
        assert not journal.is_reverted(e1.op_id)
    finally:
        server.shutdown()
        server.server_close()


def test_non_docx_doc_path_keeps_op_list_redline(tmp_path: Path) -> None:
    # A non-.docx doc must keep today's op-by-op redline behavior.
    journal, _, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0, doc_path=str(tmp_path / "notes.txt"))
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        status, body = _get(f"http://{HOST}:{port}/")
        assert "from active events" in body.decode("utf-8")  # op-list heading
        status, body = _get(f"http://{HOST}:{port}/api/redline")
        html = json.loads(body)["html"]
        assert "swift" in html  # the op-list redline still renders the change
    finally:
        server.shutdown()
        server.server_close()


def test_bad_review_request_returns_400(tmp_path: Path) -> None:
    journal, _, _ = _journal_with_ops(tmp_path / "s.changex")
    server = build_server(journal, port=0)
    port = server.server_address[1]
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    base = f"http://{HOST}:{port}"
    try:
        req = urllib.request.Request(
            base + "/api/review",
            data=json.dumps({"op_id": "x", "action": "bogus"}).encode(),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            urllib.request.urlopen(req)  # noqa: S310
            raised = False
        except urllib.error.HTTPError as exc:
            raised = exc.code == 400
        assert raised
    finally:
        server.shutdown()
        server.server_close()
